import streamlit as st
import pandas as pd
from docx import Document
from pypdf import PdfReader, PdfWriter
from pdf2docx import Converter
from reportlab.pdfgen import canvas
import io
import difflib

# --- BRANDING ---
st.set_page_config(page_title="Tamanna | File Suite", layout="wide")
st.markdown("<h1 style='color: #00f2ff;'>TAMANNA PORTAL</h1>", unsafe_allow_html=True)

menu = st.sidebar.selectbox("Select Service", ["Comparison Engine", "PDF Converter", "Split & Merge"])

# --- 1. COMPARISON ENGINE (With PDF Output) ---
if menu == "Comparison Engine":
    st.subheader("🔍 Compare & Generate Highlighted PDF")
    
    col1, col2 = st.columns(2)
    with col1:
        f1 = st.file_uploader("Upload Original (Word/Excel)", type=['docx', 'xlsx'])
    with col2:
        f2 = st.file_uploader("Upload Revised (Word/Excel)", type=['docx', 'xlsx'])
    
    if f1 and f2:
        if st.button("Submit for Comparison"):
            # Logic for Word Comparison
            if f1.name.endswith('.docx'):
                doc1 = Document(f1)
                doc2 = Document(f2)
                t1 = [p.text for p in doc1.paragraphs]
                t2 = [p.text for p in doc2.paragraphs]
                
                diff = list(difflib.ndiff(t1, t2))
                
                # Create Highlighted PDF Report
                buffer = io.BytesIO()
                p = canvas.Canvas(buffer)
                y = 800
                p.drawString(100, y, "Tamanna Comparison Report (Highlighted Changes)")
                y -= 30
                
                for line in diff:
                    if line.startswith('+ '):
                        p.setFillColorRGB(0, 0.8, 0) # Green for Added
                        p.drawString(50, y, f"[ADDED] {line[2:][:80]}")
                    elif line.startswith('- '):
                        p.setFillColorRGB(0.8, 0, 0) # Red for Removed
                        p.drawString(50, y, f"[REMOVED] {line[2:][:80]}")
                    else:
                        p.setFillColorRGB(0, 0, 0) # Black for Unchanged
                        continue
                    y -= 20
                    if y < 50: p.showPage(); y = 800
                
                p.save()
                st.success("Comparison Complete!")
                st.download_button("Download Highlighted PDF Result", buffer.getvalue(), "tamanna_diff.pdf")

# --- 2. PDF CONVERTER ---
elif menu == "PDF Converter":
    st.subheader("📄 PDF to Word")
    pdf_input = st.file_uploader("Choose PDF", type="pdf")
    if pdf_input:
        if st.button("Submit for Conversion"):
            with open("temp.pdf", "wb") as f:
                f.write(pdf_input.read())
            cv = Converter("temp.pdf")
            cv.convert("temp.docx")
            cv.close()
            with open("temp.docx", "rb") as f:
                st.download_button("Download Word File", f, "converted.docx")

# --- 3. SPLIT & MERGE ---
elif menu == "Split & Merge":
    st.subheader("📂 Split or Merge PDFs")
    mode = st.radio("Action", ["Merge", "Split"])
    
    if mode == "Merge":
        files = st.file_uploader("Upload PDFs", type="pdf", accept_multiple_files=True)
        if files and st.button("Submit Merge"):
            writer = PdfWriter()
            for f in files: writer.append(f)
            out = io.BytesIO()
            writer.write(out)
            st.download_button("Download Merged PDF", out.getvalue(), "merged.pdf")
            
    elif mode == "Split":
        s_file = st.file_uploader("Upload PDF", type="pdf")
        pg = st.number_input("Page number to extract", min_value=1, step=1)
        if s_file and st.button("Submit Split"):
            reader = PdfReader(s_file)
            writer = PdfWriter()
            writer.add_page(reader.pages[pg-1])
            out = io.BytesIO()
            writer.write(out)
            st.download_button(f"Download Page {pg}", out.getvalue(), "split.pdf")
